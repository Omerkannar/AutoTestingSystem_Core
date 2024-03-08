import datetime
import os
import subprocess
#
# import Util
from tkinter import messagebox, Image
from docx import Document
from docx.shared import Cm, Inches
import Global_Setting_Var
from PIL import Image


# input
#   strList - the name of the folder
#   type if its ATR or ATP
# does
#   make combination between  name of the folder and the first line in the ATR\ATP
# return
#   heading - string that will be the header of the page and or outline
def getHeading(strList,type):
    heading = ""

    # create the ATP\ATR full path
    if type == "ATR":
        SourcefileNameATP = Global_Setting_Var.ParentDirResu + strList + "/" + strList + "_ATR.txt"
    elif type == "ATP":
        SourcefileNameATP = Global_Setting_Var.ParentDirTest + strList + "/" + strList + "_ATP.txt"
    file = open(SourcefileNameATP, "r")
    fileLine = file.readlines()

    heading  = strList +" -> " + fileLine[1][:-1]
    return heading


# input
#   Doc - pointer for the report
#   image_path  - full image path
#   caption  - the label below the image
# does
#   insert image to the report and put caption below it


def add_image_to_docx(Doc, image_path, caption,W,H):
    img = Image.open(image_path)
    img.thumbnail((Inches(H), Inches(W)))  # Resize to a maximum size of 3x3 inches



    # Add the caption as a paragraph after the image
    para = Doc.add_paragraph(caption, style='Caption')
    para.paragraph_format.space_before=Inches(0.1)
    para.paragraph_format.space_after = Inches(0.1)
    # Add the image to the document
    Doc.add_picture(image_path, width=Inches(W), height=Inches(H))
    # Add an empty paragraph to create space between images and captions
    #Doc.add_paragraph()  # Add an empty paragraph to create space between images



def combine_images(image1_path):
    Diff_image_path = image1_path[:-4] + "_N_Diff.jpg"
    output_path = image1_path[:-7] + "_Combine.jpg"
    # Open the images
    image1 = Image.open(image1_path)
    image2 = Image.open(Diff_image_path)

    # Ensure that both images have the same height
    height = max(image1.height, image2.height)
    image1 = image1.resize((image1.width, height))
    image2 = image2.resize((image2.width, height))

    # Create a new image with combined width
    combined_width = image1.width + image2.width
    combined_image = Image.new('RGB', (combined_width, height))

    # Paste the images one below the other
    combined_image.paste(image1, (0, 0))
    #combined_image.paste(image2, (0, image1.height))
    combined_image.paste(image2, (image1.width, 0))

    # Save the result
    combined_image.save(output_path)
    return output_path

# input :
# List  - the name List of all the tests \ reports fo the report
# type  -  ATP or ATR
# SaveName  - the name of the report

def openfile(file_path):
    try:
        # Check if the 'winword' executable (Microsoft Word) is available in the system
        subprocess.run([Global_Setting_Var.wordAppPath, '/version'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        # If 'winword' is available, open the document
        subprocess.Popen([Global_Setting_Var.wordAppPath, file_path], shell=True)

    except subprocess.CalledProcessError:
        messagebox.showinfo("Error", "Microsoft Word is not installed or not found on this system.")
    except FileNotFoundError:
        messagebox.showinfo("Error", "Microsoft Word executable 'winword' not found.")
    except Exception as e:
        messagebox.showinfo("Error", f"An error occurred: {e}")

def open_folder_in_explorer(folder_path):
    try:
        subprocess.Popen(['explorer', folder_path], shell=True)
    except Exception as e:
        messagebox.showinfo("Error", f"An error occurred: {e}")


def createDocx(List,type,saveName):

    selected_indices = List.curselection()
    selected = ",".join([List.get(i).split(" ")[-1] for i in selected_indices])

    # check the user enter any test or report to the final report
    if selected == "":
        messagebox.showinfo('Auto Test', 'you need to define at least one test')
        return

    strList = selected.split(',')
    Doc = Document()
    section = Doc.sections[0]
    header = section.header
    headPargra = header.paragraphs[0]
    today = datetime.date.today()
    d2 = today.strftime("%B %d, %Y")

    # create page header
    if type == "ATR":
        headPargra.text = d2 + "\tATR for ATH Program\tAutoTest"
        numofcolum = 4
        Col2ignore1 = 3
        Col2ignore2 = 4
    elif type == "ATP":
        headPargra.text = d2 + "\tATP for ATH Program\tAutoTest"
        numofcolum = 3
        Col2ignore1 = 3
        Col2ignore2 = 100
    else:
        headPargra.text = d2 + "\tATH Program\tAutoTest"

    # create first page
    if type == "ATR":
        Doc.add_heading("\n\n\nATR for ATH  - " + d2, 0)
    else:
        Doc.add_heading("\n\n\nATP for ATH  - " + d2, 0)
    Doc.add_page_break()

    # create List of ATPs\ATRs page = outlines
    Doc.add_heading(" List Of Tests", level=1)
    for i in range(len(strList)):
        FileHeading = getHeading(strList[i],type)
        Doc.add_heading(FileHeading + "  \n", level=2)
    Doc.add_page_break()

    # for each test create its own page
    for i in range(len(strList)):
        FileHeading = getHeading(strList[i], type)
        Doc.add_heading(FileHeading + "  \n", level=2)
        pargra = Doc.add_paragraph()

        if type == "ATR":
            SourcefileNameATP = Global_Setting_Var.ParentDirResu + strList[i] + "/" + strList[i] + "_ATR.txt"
        elif type == "ATP":
            SourcefileNameATP = Global_Setting_Var.ParentDirTest + strList[i] + "/" + strList[i] + "_ATP.txt"
        else:
            messagebox.showinfo('Auto Test', 'could not create doc')
            exit()

        file = open(SourcefileNameATP, "r")
        fileLine = file.readlines()

        if type == "ATR":
            numOfline = len(fileLine)-5 # without the last  5 lines of the ATR - mouse counter
        else:
            numOfline = len(fileLine)

        lineIndex = 0
        sourceTableIndex = 0
        targTableIndex = 0
        widths = [Cm(1), Cm(6), Cm(6), Cm(3), Cm(3), Cm(3), Cm(3), Cm(3), Cm(3)]
        addTable = True
        lineIndex = 4

        #enter the ATP\ATR data to the report
        for lineIndex in range(numOfline):
            lineContainList = fileLine[lineIndex].split("\t")

            # if lineContainList > 3 its the table part
            if len(lineContainList) > 3:
                tableNumLine = numOfline - lineIndex
                if addTable == True:
                    table = Doc.add_table(rows=0, cols=numofcolum)
                    addTable = False

                table.style = 'TableGrid'
                row = table.add_row()

                # Col2ignore1 =the full image name  , Col2ignore2 = the difference value
                for tableIndex in range(len(lineContainList)):
                    if tableIndex != Col2ignore1 and tableIndex != Col2ignore2:
                        row.cells[targTableIndex].text = lineContainList[sourceTableIndex]
                        row.cells[targTableIndex].width = widths[targTableIndex]
                        targTableIndex = targTableIndex + 1
                    sourceTableIndex = sourceTableIndex + 1
                lineIndex = lineIndex + 1
                sourceTableIndex = 0
                targTableIndex = 0

            else:
                pargra.add_run(fileLine[lineIndex])
                lineIndex = lineIndex + 1

        Splitfilename = SourcefileNameATP.split("/")[:-1]
        foldername = ""
        for i in range(len(Splitfilename)):
            foldername = foldername + Splitfilename[i] + "/"

        # Folder containing the JPEG images you want to resize and add to the DOCX file
        input_folder = foldername

        # List of image file extensions
        image_extensions = (".jpg", ".jpeg")
        if Global_Setting_Var.PicOption != 0 :
            for filename in os.listdir(input_folder):
                if filename.lower().endswith(image_extensions) and not("Diff" in filename) and not("Combine" in filename) and not("concatenate" in filename)and not("cut" in filename)and not("GS" in filename):
                    image_path = os.path.join(input_folder, filename)
                    if type == "ATR" and Global_Setting_Var.PicOption == 2:
                        image_path = combine_images(image_path)

                    if Global_Setting_Var.PicOption == 1:
                        Width = 3
                        Height = 3
                    else:
                        Width = 6
                        Height = 2.8


                    add_image_to_docx(Doc, image_path, filename, Width, Height)

            Doc.add_page_break()
    try:
        Doc.save(saveName)
        open_folder_in_explorer(Global_Setting_Var.docFolder)
        # openfile(Global_Setting_Var.docFolder +"\\" +saveName.split("\\")[-1])
    except Exception as e:
        messagebox.showinfo("Error", f"An error occurred: {e}")



